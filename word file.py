import pandas as pd
import os
from docx import Document

# Paths
folder = r'D:\Axon python scripts\Documentation'
diff_file = os.path.join(folder, 'diff_analysis.xlsx')
output_file = os.path.join(folder, 'UST_Differential_Fieldwise.docx')

# Read excel
df = pd.read_excel(diff_file)

# Updated: Fields to show, including LUST fields
field_tuples = [
    ('Facility Count','Facility Count_Current','Facility Count_Previous','Facility Count_Diff'),
    ('Tank Count','Tank Count_Current','Tank Count_Previous','Tank Count_Diff'),
    ('City Count','City Count_Current','City Count_Previous','City Count_Diff'),
    ('Zip Count','Zip Count_Current','Zip Count_Previous','Zip Count_Diff'),
    ('UST/AST Types','UST/AST Types_Current','UST/AST Types_Previous','UST/AST Types_Diff'),
    ('Install Years','Install Years_Current','Install Years_Previous','Install Years_Diff'),
    ('Tank Size Min','Tank Size Min_Current','Tank Size Min_Previous','Tank Size Min_Diff'),
    ('Tank Size Max','Tank Size Max_Current','Tank Size Max_Previous','Tank Size Max_Diff'),
    ('Tank Size Avg','Tank Size Avg_Current','Tank Size Avg_Previous','Tank Size Avg_Diff'),
    ('Top Size Range','Top Size Range_Current','Top Size Range_Previous','Top Size Range_Changed'),
    ('Top Size Range Count','Top Size Range Count_Current','Top Size Range Count_Previous','Top Size Range Count_Diff'),
    ('Top Tank Status','Top Tank Status_Current','Top Tank Status_Previous','Top Tank Status_Changed'),
    ('Top Tank Status Count','Top Tank Status Count_Current','Top Tank Status Count_Previous','Top Tank Status Count_Diff'),
    ('Top Construction','Top Construction_Current','Top Construction_Previous','Top Construction_Changed'),
    ('Top Construction Count','Top Construction Count_Current','Top Construction Count_Previous','Top Construction Count_Diff'),
    ('Top Content','Top Content_Current','Top Content_Previous','Top Content_Changed'),
    ('Top Content Count','Top Content Count_Current','Top Content Count_Previous','Top Content Count_Diff'),
    ('Top Piping Construction','Top Piping Construction_Current','Top Piping Construction_Previous','Top Piping Construction_Changed'),
    ('Top Piping Const Count','Top Piping Const Count_Current','Top Piping Const Count_Previous','Top Piping Const Count_Diff'),
    ('County Count','County Count_Current','County Count_Previous','County Count_Diff'),
    ('Last Sync Min','Last Sync Min_Current','Last Sync Min_Previous',''),
    ('Last Sync Max','Last Sync Max_Current','Last Sync Max_Previous',''),
    # --- LUST terms, added below ---
    ('Top Lust Term', 'Top Lust Term_Current', 'Top Lust Term_Previous', 'Top Lust Term_Changed'),
    ('Top Lust Count', 'Top Lust Count_Current', 'Top Lust Count_Previous', 'Top Lust Count_Diff'),
]

doc = Document()
doc.add_heading('UST Statewise Differential Field Analysis Report', 0)

for idx, row in df.iterrows():
    state = str(row['State'])
    doc.add_heading(state, level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current'
    hdr_cells[2].text = 'Previous'
    hdr_cells[3].text = 'Difference'

    for label, cur_col, prev_col, diff_col in field_tuples:
        cur_val = row.get(cur_col, "")
        prev_val = row.get(prev_col, "")
        if diff_col:
            diff_val = row.get(diff_col, "")
        else:
            diff_val = "0" if cur_val == prev_val else "-"

        # Show 'Changed' for boolean flag columns
        if diff_col and 'Changed' in diff_col:
            if str(diff_val).strip().upper() in ['TRUE', '1']:
                diff_val = "Changed"
            else:
                diff_val = "0"

        # Clean NaNs
        cur_val = "" if pd.isna(cur_val) else str(cur_val)
        prev_val = "" if pd.isna(prev_val) else str(prev_val)

        # Optional: Format floats differently
        try:
            if isinstance(float(cur_val), float) and cur_val != '' and abs(float(cur_val)) > 1000:
                cur_val = f"{float(cur_val):,.2f}"
        except:
            pass
        try:
            if isinstance(float(prev_val), float) and prev_val != '' and abs(float(prev_val)) > 1000:
                prev_val = f"{float(prev_val):,.2f}"
        except:
            pass
        try:
            if diff_val not in ["Changed", "0", "-"]:
                diff_val = f"{float(diff_val):,.2f}" if abs(float(diff_val)) > 1000 else str(diff_val)
        except:
            pass

        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = cur_val
        row_cells[2].text = prev_val
        row_cells[3].text = str(diff_val)
    doc.add_paragraph("")

doc.save(output_file)
print(f"\nWord report saved at: {output_file}")